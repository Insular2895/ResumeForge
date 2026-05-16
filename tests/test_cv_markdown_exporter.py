from pathlib import Path
import json

from docx import Document

from src.application.cv_markdown_exporter import export_cv_markdown


def test_cv_markdown_exporter_preserves_report_and_docx_content(tmp_path):
    docx_path = tmp_path / "CV_Test.docx"
    document = Document()
    document.add_paragraph("Lucas Pertusa")
    document.add_paragraph("Coordination de flux internationaux avec SAP et 12 partenaires")
    document.save(docx_path)

    report_path = tmp_path / "last_run_report.json"
    report_path.write_text(
        json.dumps(
            {
                "company_detected": "Ipsen",
                "job_title_detected": "Coordinateur ADV Import-Export",
                "selected_experiences": [
                    {
                        "company": "Blurry",
                        "position_title": "Operations",
                        "dates": "2024",
                        "reason_tags": ["SAP", "stock"],
                    }
                ],
                "selected_certifications": ["SAP Supply Chain"],
                "selected_technical_skills": ["SAP", "Import-export"],
            }
        ),
        encoding="utf-8",
    )

    output_path = export_cv_markdown(docx_path, report_path)
    markdown = Path(output_path).read_text(encoding="utf-8")

    assert "Ipsen" in markdown
    assert "Coordinateur ADV Import-Export" in markdown
    assert "Blurry" in markdown
    assert "SAP" in markdown
    assert "12 partenaires" in markdown

