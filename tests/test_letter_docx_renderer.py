from docx import Document

from src.letter.letter_docx_renderer import render_letter_docx


def test_letter_docx_renderer_replaces_placeholders(tmp_path):
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Paris, le [[LM_DATE]]")
    document.add_paragraph("Objet : [[LM_JOB_TITLE]] chez [[LM_COMPANY]]")
    document.add_paragraph("[[LM_FINAL_LETTER]]")
    document.add_paragraph("[[LM_SIGNATURE]]")
    document.save(template_path)

    output_path = tmp_path / "LM.docx"
    render_letter_docx(
        {"company": "Ipsen", "job_title": "Coordinateur ADV Import-Export"},
        "Madame, Monsieur, je vous propose ma candidature.",
        output_path,
        template_path=template_path,
    )

    rendered = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "[[" not in text
    assert "Ipsen" in text
    assert "Coordinateur ADV Import-Export" in text

