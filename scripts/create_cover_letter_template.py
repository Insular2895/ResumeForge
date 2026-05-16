from pathlib import Path

from docx import Document


ROOT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT_DIR / "templates" / "base_cover_letter_example.docx"


def create_template(output_path: Path = OUTPUT_PATH) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_paragraph("[[CANDIDATE_FULL_NAME]]")
    document.add_paragraph("[[CANDIDATE_ADDRESS_LINE_1]]")
    document.add_paragraph("[[CANDIDATE_ADDRESS_LINE_2]]")
    document.add_paragraph("tel : [[CANDIDATE_PHONE]]")
    document.add_paragraph("email : [[CANDIDATE_EMAIL]]")
    document.add_paragraph("")
    document.add_paragraph("a        [[LM_COMPANY]]")
    document.add_paragraph("[[LM_COMPANY_ADDRESS_LINE_1]]")
    document.add_paragraph("[[LM_COMPANY_POSTAL_CITY]]")
    document.add_paragraph("")
    document.add_paragraph("A l'attention de [[LM_ATTENTION_TO]]")
    document.add_paragraph("[[LM_DEPARTMENT]]")
    document.add_paragraph("")
    document.add_paragraph("[[LM_SALUTATION]],")
    document.add_paragraph("[[LM_FINAL_LETTER]]")
    document.add_paragraph("")
    document.add_paragraph("Le [[LM_DATE]]")
    document.add_paragraph("[[LM_SIGNATURE]]")
    document.save(output_path)
    return output_path


def main() -> None:
    path = create_template()
    print(f"Template exemple cree : {path}")


if __name__ == "__main__":
    main()
