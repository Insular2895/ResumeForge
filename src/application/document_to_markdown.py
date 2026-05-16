from pathlib import Path

from typing import Protocol

from docx import Document


class OptionalDocxMarkdownConverter(Protocol):
    """Optional adapter interface for tools such as MinerU or Marker."""

    def convert(self, docx_path: Path) -> str:
        ...


def docx_to_markdown(
    docx_path: str | Path,
    optional_converter: OptionalDocxMarkdownConverter | None = None,
) -> str:
    """Convert a DOCX document to a conservative Markdown text fallback."""
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"CV DOCX not found: {path}")

    if optional_converter is not None:
        converted = optional_converter.convert(path)
        if converted.strip():
            return converted.strip() + "\n"

    document = Document(path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if "heading" in style_name or "titre" in style_name:
            lines.append(f"## {text}")
        elif style_name.startswith("list") or text.startswith(("•", "-", "*")):
            lines.append(f"- {text.lstrip('•-* ').strip()}")
        else:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n\n".join(lines).strip() + "\n"
