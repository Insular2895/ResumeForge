from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DEFAULT_FONT_NAME = "Arial"


class DocxTemplateRenderer:
    def __init__(self, template_path):
        self.template_path = Path(template_path)

        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")

    def render(self, replacements, output_path):
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document(self.template_path)

        self._force_document_font(doc, DEFAULT_FONT_NAME)

        for paragraph in list(doc.paragraphs):
            self._replace_paragraph(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in list(cell.paragraphs):
                        self._replace_paragraph(paragraph, replacements)

        self._force_document_font(doc, DEFAULT_FONT_NAME)
        self._force_font_sizes(doc)

        doc.save(output_path)
        return output_path

    def _replace_paragraph(self, paragraph, replacements):
        paragraph_text = paragraph.text

        if not paragraph_text:
            return

        bullet_placeholders = {
            "[[EXP_1_BULLETS]]",
            "[[EXP_2_BULLETS]]",
            "[[LEAD_1_BULLETS]]",
            "[[CERTIFICATION_ENTRIES]]",
        }

        paragraph_placeholders = {
            "[[LM_FINAL_LETTER]]",
        }

        for placeholder in bullet_placeholders:
            if placeholder in paragraph_text:
                bullets_text = replacements.get(placeholder, "")
                bullets = self._split_bullets(bullets_text)
                self._replace_bullet_placeholder_with_paragraphs(paragraph, bullets)
                return

        for placeholder in paragraph_placeholders:
            if placeholder in paragraph_text:
                paragraph_text = replacements.get(placeholder, "")
                paragraphs = self._split_text_paragraphs(paragraph_text)
                self._replace_text_placeholder_with_paragraphs(paragraph, paragraphs)
                return

        self._replace_inline_placeholders(paragraph, replacements)

    def _replace_inline_placeholders(self, paragraph, replacements):
        """
        Remplace les placeholders simples sans casser le gras existant.
        Si un placeholder est en gras dans Word, la valeur remplacée reste en gras.
        """
        original_text = paragraph.text

        if not original_text:
            return

        label_prefixes = [
            "Compétences techniques :",
            "Intérêts :",
            "Langues :",
            "Certifications :",
            "Formation :",
        ]

        full_text = original_text
        for key, value in replacements.items():
            full_text = full_text.replace(key, str(value))

        for label in label_prefixes:
            if full_text.startswith(label):
                self._set_label_line(paragraph, full_text, label)
                return

        # Remplacement run par run = conserve gras, italique, taille, etc.
        for run in paragraph.runs:
            run_text = run.text

            for key, value in replacements.items():
                if key in run_text:
                    run_text = run_text.replace(key, str(value))

            run.text = run_text
            self._force_run_font(run, DEFAULT_FONT_NAME)

        # Fallback si Word a coupé un placeholder entre plusieurs runs
        remaining_text = paragraph.text
        if any(key in remaining_text for key in replacements.keys()):
            self._set_paragraph_text_preserve_style(paragraph, full_text)

    def _replace_bullet_placeholder_with_paragraphs(self, paragraph, bullets):
        if not bullets:
            self._remove_paragraph(paragraph)
            return

        self._set_paragraph_text_preserve_style(paragraph, bullets[0])

        previous = paragraph

        for bullet in bullets[1:]:
            new_paragraph = self._insert_paragraph_after(previous)
            self._copy_paragraph_format(paragraph, new_paragraph)
            self._set_paragraph_text_preserve_style(new_paragraph, bullet)
            previous = new_paragraph

    def _split_bullets(self, bullets_text):
        if bullets_text is None:
            return []

        if isinstance(bullets_text, list):
            raw_lines = bullets_text
        else:
            raw_lines = str(bullets_text).splitlines()

        bullets = []

        for line in raw_lines:
            line = str(line).strip()

            if not line:
                continue

            line = line.lstrip("•").strip()
            line = line.lstrip("-").strip()

            if line:
                bullets.append(line)

        return bullets

    def _split_text_paragraphs(self, text):
        if text is None:
            return []

        raw = str(text).replace("\r\n", "\n").replace("\r", "\n").strip()
        if not raw:
            return []

        parts = re.split(r"\n\s*\n+", raw)
        return [re.sub(r"\s*\n\s*", " ", part).strip() for part in parts if part.strip()]

    def _replace_text_placeholder_with_paragraphs(self, paragraph, paragraphs):
        if not paragraphs:
            self._remove_paragraph(paragraph)
            return

        self._set_paragraph_text_preserve_style(paragraph, paragraphs[0])
        paragraph.paragraph_format.space_after = Pt(8)

        previous = paragraph
        for text in paragraphs[1:]:
            new_paragraph = self._insert_paragraph_after(previous)
            self._copy_paragraph_format(paragraph, new_paragraph)
            new_paragraph.paragraph_format.space_after = Pt(8)
            self._set_paragraph_text_preserve_style(new_paragraph, text)
            previous = new_paragraph

    def _set_label_line(self, paragraph, text, label):
        """
        Pour :
        Compétences techniques :
        Intérêts :
        Langues :

        Seul le label est en gras.
        """
        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            title_run = paragraph.runs[0]
        else:
            title_run = paragraph.add_run()

        title_run.text = label
        title_run.bold = True
        self._force_run_font(title_run, DEFAULT_FONT_NAME)

        content = text.replace(label, "", 1)

        content_run = paragraph.add_run(content)
        content_run.bold = False
        self._force_run_font(content_run, DEFAULT_FONT_NAME)

    def _set_paragraph_text_preserve_style(self, paragraph, text):
        text = str(text)

        if paragraph.runs:
            first_run = paragraph.runs[0]

            for run in paragraph.runs:
                run.text = ""

            first_run.text = text
            self._force_run_font(first_run, DEFAULT_FONT_NAME)
        else:
            run = paragraph.add_run(text)
            self._force_run_font(run, DEFAULT_FONT_NAME)

        for run in paragraph.runs:
            self._force_run_font(run, DEFAULT_FONT_NAME)

    def _remove_paragraph(self, paragraph):
        try:
            p = paragraph._p
            p.getparent().remove(p)
        except Exception:
            self._clear_paragraph(paragraph)

    def _clear_paragraph(self, paragraph):
        for run in paragraph.runs:
            run.text = ""
            self._force_run_font(run, DEFAULT_FONT_NAME)

    def _insert_paragraph_after(self, paragraph):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    def _copy_paragraph_format(self, source, target):
        target.style = source.style

        target_format = target.paragraph_format
        source_format = source.paragraph_format

        target_format.left_indent = source_format.left_indent
        target_format.right_indent = source_format.right_indent
        target_format.first_line_indent = source_format.first_line_indent
        target_format.space_before = source_format.space_before
        target_format.space_after = source_format.space_after
        target_format.line_spacing = source_format.line_spacing
        target_format.alignment = source_format.alignment

        source_ppr = source._p.get_or_add_pPr()
        target_ppr = target._p.get_or_add_pPr()

        source_numpr = source_ppr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
        )

        if source_numpr is not None:
            target_ppr.append(deepcopy(source_numpr))

    def _force_document_font(self, doc, font_name):
        for style in doc.styles:
            try:
                if hasattr(style, "font"):
                    style.font.name = font_name

                    if style._element.rPr is not None:
                        r_fonts = style._element.rPr.rFonts

                        if r_fonts is not None:
                            r_fonts.set(qn("w:ascii"), font_name)
                            r_fonts.set(qn("w:hAnsi"), font_name)
                            r_fonts.set(qn("w:eastAsia"), font_name)
                            r_fonts.set(qn("w:cs"), font_name)
            except Exception:
                pass

        for paragraph in doc.paragraphs:
            self._force_paragraph_font(paragraph, font_name)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._force_paragraph_font(paragraph, font_name)

    def _force_paragraph_font(self, paragraph, font_name):
        for run in paragraph.runs:
            self._force_run_font(run, font_name)

    def _force_run_font(self, run, font_name):
        run.font.name = font_name

        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts

        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)

        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:cs"), font_name)

    def _force_font_sizes(self, doc):
        """
        Force 10pt partout, 11pt pour les paragraphes dont au moins un run est bold.
        Évite le mélange de tailles causé par l'injection de runs via python-docx.
        """
        all_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for paragraph in all_paragraphs:
            if not paragraph.runs:
                continue

            is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
            target_size = Pt(11) if is_bold else Pt(10)

            for run in paragraph.runs:
                run.font.size = target_size
